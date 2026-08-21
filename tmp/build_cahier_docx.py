from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
MARKDOWN_PATH = ROOT / "CAHIER_DES_CHARGES_FONCTIONNALITES_EXISTANTES.md"
OUTPUT_PATH = ROOT / "Cahier de charge actualise - Walikale Papeterie.docx"
LOGO_PATH = ROOT / "img" / "logo-walikale1.png"


BLUE = RGBColor(19, 86, 210)
NAVY = RGBColor(10, 31, 68)
TEXT = RGBColor(33, 43, 54)
MUTED = RGBColor(97, 113, 132)
BORDER = "D9E3F0"
LIGHT = "F4F8FE"


def set_font(run, name="Arial", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_borders(table, color=BORDER, size="8"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_padding(cell, top=80, bottom=80, start=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in {"top": top, "bottom": bottom, "start": start, "end": end}.items():
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def clear_paragraph(paragraph):
    p = paragraph._element
    for child in list(p):
        p.remove(child)


def configure_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(11)
    normal.font.color.rgb = TEXT
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, color, before, after in (
        ("Heading 1", 16, NAVY, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, NAVY, 10, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.1

    if "Body Small" not in doc.styles:
        small = doc.styles.add_style("Body Small", WD_STYLE_TYPE.PARAGRAPH)
        small.base_style = normal
        small.font.name = "Arial"
        small._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        small._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        small.font.size = Pt(9.5)
        small.font.color.rgb = MUTED
        small.paragraph_format.space_after = Pt(4)
        small.paragraph_format.line_spacing = 1.1


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if LOGO_PATH.exists():
        p.add_run().add_picture(str(LOGO_PATH), width=Inches(1.2))
    p.paragraph_format.space_after = Pt(10)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(4)
    r = kicker.add_run("CAHIER DE CHARGE ACTUALISE")
    set_font(r, size=11, color=BLUE, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    r = title.add_run("Walikale Papeterie")
    set_font(r, size=24, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(16)
    r = subtitle.add_run("Version de reference des fonctionnalites existantes")
    set_font(r, size=12, color=MUTED)

    meta = doc.add_table(rows=3, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.columns[0].width = Inches(2.0)
    meta.columns[1].width = Inches(4.2)
    set_table_borders(meta)
    rows = [
        ("Structure", "Walikale to World"),
        ("Date de mise a jour", "21 aout 2026"),
        ("Objet", "Validation de l'existant avant deploiement"),
    ]
    for row, (label, value) in zip(meta.rows, rows):
        row.cells[0].width = Inches(2.0)
        row.cells[1].width = Inches(4.2)
        set_cell_shading(row.cells[0], LIGHT)
        for cell in row.cells:
            set_cell_padding(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        clear_paragraph(row.cells[0].paragraphs[0])
        clear_paragraph(row.cells[1].paragraphs[0])
        p1 = row.cells[0].paragraphs[0]
        p2 = row.cells[1].paragraphs[0]
        rr1 = p1.add_run(label)
        rr2 = p2.add_run(value)
        set_font(rr1, size=10.5, color=NAVY, bold=True)
        set_font(rr2, size=10.5, color=TEXT)

    doc.add_paragraph("")

    intro = doc.add_paragraph()
    intro.paragraph_format.space_after = Pt(10)
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = intro.add_run(
        "Ce document presente, de maniere professionnelle et consolidée, l'ensemble des fonctionnalites actuellement actives "
        "dans l'application Walikale Papeterie, en distinguant les usages desktop, web, hors ligne et cloud."
    )
    set_font(r, size=11, color=TEXT)

    summary_title = doc.add_paragraph()
    summary_title.paragraph_format.space_before = Pt(6)
    summary_title.paragraph_format.space_after = Pt(6)
    r = summary_title.add_run("Synthese executive")
    set_font(r, size=13, color=NAVY, bold=True)

    summary_box = doc.add_table(rows=1, cols=1)
    summary_box.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(summary_box, color="CFE0F5")
    cell = summary_box.cell(0, 0)
    set_cell_shading(cell, "F7FAFF")
    set_cell_padding(cell, top=120, bottom=120, start=160, end=160)
    bullets = [
        "Application de gestion de papeterie, stock, ventes, services, depenses et utilisateurs.",
        "Fonctionnement hybride : desktop Windows hors ligne avec SQLite et web en ligne avec Supabase.",
        "Synchronisation selective par rubrique entre le poste local et le cloud.",
        "Rapports financiers, factures, tickets, historique d'activite et sauvegardes.",
    ]
    clear_paragraph(cell.paragraphs[0])
    for index, bullet in enumerate(bullets):
        p = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        p.style = doc.styles["Normal"]
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.first_line_indent = Inches(-0.15)
        run = p.add_run(f"• {bullet}")
        set_font(run, size=11, color=TEXT)

    doc.add_page_break()


def add_header_footer(doc):
    section = doc.sections[0]
    header = section.header
    hp = header.paragraphs[0]
    clear_paragraph(hp)
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r1 = hp.add_run("Walikale Papeterie")
    set_font(r1, size=9, color=NAVY, bold=True)
    r2 = hp.add_run("  |  Cahier de charge actualise")
    set_font(r2, size=9, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    clear_paragraph(fp)
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = fp.add_run("Document interne de reference")
    set_font(r, size=9, color=MUTED)


def parse_markdown(markdown_text):
    lines = markdown_text.splitlines()
    blocks = []
    current = None
    for raw in lines:
        line = raw.rstrip()
        if not line.strip():
            continue
        if line.startswith("# "):
            blocks.append(("h0", line[2:].strip()))
        elif line.startswith("## "):
            blocks.append(("h1", line[3:].strip()))
        elif line.startswith("### "):
            blocks.append(("h2", line[4:].strip()))
        elif line.startswith("- "):
            blocks.append(("bullet", line[2:].strip()))
        else:
            blocks.append(("p", line.strip()))
    return blocks


def add_content(doc, blocks):
    skip_titles = {"Cahier Des Charges", "Application", "Objet du document"}
    skip_values = {"Walikale Papeterie"}
    first_body = True
    for kind, text in blocks:
        if text in skip_titles or text in skip_values:
            continue
        if kind == "h0":
            continue
        if kind == "h1":
            p = doc.add_paragraph(style="Heading 1")
            p.paragraph_format.keep_with_next = True
            run = p.add_run(text)
            set_font(run, size=16, color=NAVY, bold=True)
        elif kind == "h2":
            p = doc.add_paragraph(style="Heading 2")
            p.paragraph_format.keep_with_next = True
            run = p.add_run(text)
            set_font(run, size=13, color=BLUE, bold=True)
        elif kind == "bullet":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.first_line_indent = Inches(-0.18)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(f"• {text}")
            set_font(run, size=11, color=TEXT)
        else:
            if first_body and text.startswith("Ce document"):
                first_body = False
                continue
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(6)
            clean = re.sub(r"`([^`]+)`", r"\1", text)
            run = p.add_run(clean)
            set_font(run, size=11, color=TEXT)


def build_doc():
    markdown_text = MARKDOWN_PATH.read_text(encoding="utf-8")
    doc = Document()
    configure_styles(doc)
    add_cover(doc)
    add_header_footer(doc)
    blocks = parse_markdown(markdown_text)
    add_content(doc, blocks)
    doc.save(OUTPUT_PATH)
    print(str(OUTPUT_PATH))


if __name__ == "__main__":
    build_doc()
